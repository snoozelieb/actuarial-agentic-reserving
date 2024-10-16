"""
Title: Lessons and cuntionalities I add in this version 

Human input:    Human input is needed to make tihs work.
                I follow the guidelines from the following docs:https://docs.crewai.com/how-to/Human-Input-on-Execution/
                Set human_input = True in tasks to allow for human input

Cache:          I also add cache functionality which was recommended by the crewai course.

Context:        I also add the context variable to the tasks to allow for the agents to have a better understanding of the task at hand. #https://docs.crewai.com/core-concepts/Tasks/#task-attributes
"""


#######################################################################################################
import os
import pandas as pd
import numpy as np
from scipy import stats
import matplotlib.pyplot as plt
import seaborn as sns
from crewai import Agent, Task, Crew, Process
from crewai_tools import SerperDevTool

#Word document dependencies
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

#https://www.youtube.com/watch?v=q6QLGS306d0&ab_channel=TylerAI Crew ai tutorial that the code is based on
os.environ["OPENAI_API_KEY"] = "api key here"
os.environ["OPENAI_MODEL_NAME"]= "gpt-4o"

#############################################################################################################################################################################################################
###################################################################################################################################################################################################
"""
Non-crewai functions

"""
#######################################################################################################################################
#Reading in CSV file
# Function to read CSV file
def read_csv_file(file_path):
    try:
        df = pd.read_csv(file_path)
        return df.to_json(orient='split')
    except Exception as e:
        return f"Error reading CSV file: {str(e)}"

# Get CSV file path from user
csv_file_path = input("Enter the path to your actuarial data CSV file: ")

# Read CSV file
csv_data = read_csv_file(csv_file_path)


#def export_results(results, output_dir='output'):
    # Create output directory if it doesn't exist
#    os.makedirs(output_dir, exist_ok=True)#

    # Export to text file
#    export_to_text(results, os.path.join(output_dir, 'results.txt'))
#
#    # Export to Word document
#    export_to_word(results, os.path.join(output_dir, 'results.docx'))

#def export_to_text(results, filename):
#    with open(filename, 'w') as f:
#        for section, content in results.items():
#            f.write(f"{section.upper()}\n{'=' * len(section)}\n\n")
#            if isinstance(content, dict):
#                for subsection, subcontent in content.items():
#                    f.write(f"{subsection}\n{'-' * len(subsection)}\n")
#                    f.write(f"{subcontent}\n\n")
#            else:
#                f.write(f"{content}\n\n")

#def export_to_word(results, filename):
#    doc = Document()

    # Define styles
#    styles = doc.styles
#    style_heading1 = styles.add_style('Heading1', WD_STYLE_TYPE.PARAGRAPH)
#    style_heading1.font.size = Pt(16)
#    style_heading1.font.bold = True

#    style_heading2 = styles.add_style('Heading2', WD_STYLE_TYPE.PARAGRAPH)
#    style_heading2.font.size = Pt(14)
#    style_heading2.font.bold = True

#    for section, content in results.items():
#        doc.add_paragraph(section.upper(), style='Heading1')
#        if isinstance(content, dict):
#            for subsection, subcontent in content.items():
#                doc.add_paragraph(subsection, style='Heading2')
#                doc.add_paragraph(subcontent)
#        else:
#            doc.add_paragraph(content)
#        doc.add_paragraph()  # Add an empty paragraph for spacing

#    doc.save(filename)



#############################################################################################################################################################################################################
###################################################################################################################################################################################################

"""
Crewai specific functions
"""


# Define agents
data_analyst = Agent(
    role='Data Analyst',
    goal='Process and analyze actuarial data for multiple Lines of Business.You should identify data types, assess data quality, provide insights, and guide the user through the process of analyzing and splitting data by LOB if they are multiple LOB.Ensure compliance with APN 401 Section 3 (Data) and consider user inputs for data-driven decisions',
    backstory="""You are a skilled data analyst specializing in actuarial data. 
    You excel at identifying data types, assessing data quality, and providing insights. You also excel at making run-off triangles from raw data and passing it to the development factor specialist.""",
    verbose=True,
    allow_delegation=False,
    cache = True, #enable cache feature for the agent https://docs.crewai.com/how-to/Human-Input-on-Execution/#example
    # You can pass an optional llm attribute specifying what model you wanna use.
    # llm=ChatOpenAI(model_name="gpt-3.5", temperature=0.7),
)

development_factor_specialist = Agent(
    role='Development Factor Specialist',
    goal='Review the accuracy of the run off triangles created based on input data. Use the run-off triangles to  Calculate development factors and apply appropriate reserving methods',
    backstory="""You are an experienced actuary specializing in calculating development factors for reserving. 
    You are proficient in methods like Chain Ladder only.""",
    verbose=True,
    allow_delegation=False,
    cache = True,
    
)

risk_adjustment_specialist = Agent(
    role='Risk Adjustment Specialist',
    goal='Calculate risk adjustments using bootstrapping techniques',
    backstory="""You are an expert in calculating risk adjustments for insurance reserves. 
    You specialize in using bootstrapping techniques to calculate adjustments at the 75th percentile.""",
    verbose=True,
    allow_delegation=False,
    cache = True,
    
)

report_compiler = Agent(
    role='Report Compiler',
    goal='Compile a comprehensive, regulation-compliant actuarial report',
    backstory="""You are a meticulous report compiler with extensive knowledge of actuarial standards. 
    You excel at integrating complex analyses into clear, compliant reports.""",
    verbose=True,
    allow_delegation=False,
    cache = True,
    
)
#######################################################################################################################################
# Define tasks
##Tasks are the actual work that needs to be done by the agents. So we give the agents tasks
#  Parameters for tasks = the description, expected output (How is the information going to presented...Report, dataframe, informale email...) and the agent that will be doing the task
#Crewai allows one to assign a task to a single agent or multiple agents
# Define tasks
task0 = Task(
    description=f"""Convert the claims data into the Actuarial run-off triangle by accident year based on the data, and in preparation for claims reserving. 
    The data is provided in JSON format as follows: {csv_data}.""",
    agent=data_analyst,
    expected_output="Detailed data analysis report",
    human_input=False #https://docs.crewai.com/how-to/Human-Input-on-Execution/#example #ensures the triangels made from task 1 are used in task 2
    #Human input here is key to ensure that the triangles made are in line with our expectations.
)

task1 = Task(
    description=f"""Analyze the provided actuarial data CSV file. 
    The data is provided in JSON format as follows: {csv_data}
    Create a run-off trinagle based on the claims data provided, assess data quality, Check for data errors such as paid date being earlier than reported or incurred date and produce a column summarizing all errors found, and prepare the data for further analysis. 
    Provide a summary of your findings.""",
    agent=data_analyst,
    expected_output="Detailed data analysis report",
    human_input=False #https://docs.crewai.com/how-to/Human-Input-on-Execution/#example #ensures the triangels made from task 1 are used in task 2
)

task2 = Task(
    description="""Using the prepared data from the previous task, calculate development factors and apply appropriate reserving methods. 
    Consider methods like Chain Ladder and Bornhuetter-Ferguson. 
    Provide reserve estimates at best estimate (50th percentile), prudent (75th percentile), and optimistic (25th percentile).""",
    agent=development_factor_specialist,
    expected_output="Reserving analysis report with development factors and reserve estimates",
    human_input=False, #https://docs.crewai.com/how-to/Human-Input-on-Execution/#example
    #context=[task0,task1]  #https://docs.crewai.com/core-concepts/Tasks/#task-attributes
    #editted out context as it requires too many tokens
)

task3 = Task(
    description="""Calculate risk adjustments using bootstrapping techniques at the 75th percentile. 
    Ensure alignment with IFRS 17 and APN 401 Section 6. 
    Document your methods, assumptions, and justification for margins.""",
    agent=risk_adjustment_specialist,
    expected_output="Risk adjustment analysis report"
)

task4 = Task(
    description="""Compile a comprehensive actuarial report integrating all previous analyses. 
    Include sections on data quality, methodologies, reserve calculations, and risk adjustments. 
    Ensure compliance with APN 401 and IFRS 17 standards.""",
    agent=report_compiler,
    expected_output="Final comprehensive actuarial report"
)

# Assemble the crew
crew = Crew(
    agents=[data_analyst, development_factor_specialist, risk_adjustment_specialist, report_compiler],
    tasks=[task0, task1, task2, task3, task4],
    verbose=True,
    process=Process.sequential,
    #memory=True, #enable memory feature for the crew     https://docs.crewai.com/how-to/Human-Input-on-Execution/#example
    #planning=True #enable planning feature for the crew https://docs.crewai.com/how-to/Human-Input-on-Execution/#example
    #edited out memory and planning as it lead to too many tokens
)

# Execute the analysis
result = crew.kickoff()

#Addtional functionality to export results to a text file and word document 
print("Results exported successfully!")

print("######################")
print(result)